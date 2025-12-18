from docx import Document
from pathlib import Path

def make_docx(text: str, out_path: str):
    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        for line in block.splitlines():
            doc.add_paragraph(line)
        doc.add_paragraph("")
    doc.save(str(out))
